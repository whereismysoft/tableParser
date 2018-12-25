# before run:
#     create env (pip3)
#     install package docx
# if env created run comand: source tutorial-env/bin/activate (Linux) || tutorial-env\Scripts\activate.bat (Windows)
# then run file and write file name with extension

from docx.api import Document
import simplejson as json
import sys
import re

# print("The arguments are: " , Document)

document = Document(str(sys.argv[1]))

outputDoc = Document()

keys = tuple(('id','description'))

data = []

count = 0

for i in document.paragraphs:
    value = []
    value.append(count)
    
    value.append(str(i.text))
    # below - example of python reg exp using modular 're'
    # value.append(re.sub(r'\d+\.', '', i.text)) 
    data.append(dict(zip(keys, value)))

    count = count + 1

j = json.dumps(data, ensure_ascii=False, encoding='utf8')

with open('data.json', 'w') as f:
    f.write(j)

# table = document.tables[0]

# data = []

# keys = None
# for i, row in enumerate(table.rows):
#     text = list(cell.text for cell in row.cells)

#     values = []
#     # Establish the mapping based on the first row
#     # headers; these will become the keys of our dictionary
#     if i == 0:
#         # keys = tuple(text)
#         keys = tuple(('id','building', 'names'))
#         continue
#     values.append(i)
#     values.append(text[1])
#     values.append(text[2])

#     for n, el in enumerate(values):
#         if (n == 1):
#             values[n] = el.replace('\n', '')
#         if (n == 2):
#             values[n] = el.split('\n')
        
#     # Construct a dictionary for this row, mapping
#     # keys to values for this row
#     row_data = dict(zip(keys, values))
#     data.append(row_data)

# outputDoc.add_paragraph(str(data))

# outputDoc.save('demo.docx')
